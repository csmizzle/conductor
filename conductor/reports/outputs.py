"""
Conductor Report Outputs
- PDF Report
"""
import pdfkit
from langchain.prompts import PromptTemplate
from conductor.reports.models import ParsedReport, Report, ReportV2
from conductor.llms import openai_gpt_4o
from conductor.zen import get_image
from textwrap import dedent
import tempfile
from langchain.output_parsers import PydanticOutputParser
from langchain.chains.llm import LLMChain
from langsmith import traceable
from docx import Document
from jinja2 import Environment, FileSystemLoader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ListStyle, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListItem,
    ListFlowable,
    Image,
    PageBreak,
    Table,
    TableStyle,
)
import html
import os
from PIL import UnidentifiedImageError
import logging

BASEDIR = os.path.dirname(os.path.abspath(__file__))
logger = logging.getLogger(__name__)


def report_to_html(report: Report) -> str:
    """
    Convert a report to an HTML string
    """
    env = Environment(loader=FileSystemLoader(os.path.join(BASEDIR, "templates")))
    report_template = env.get_template("report.html")
    output_from_parsed_template = report_template.render(
        report_title=report.report.title,
        report_description=report.report.description,
        report_sections=[section.dict() for section in report.report.sections],
    )
    return html.unescape(output_from_parsed_template)


def report_v2_to_html(report: ReportV2, graph_file: str = None) -> str:
    """Generate a report with a graph embedded in the HTML

    Args:
        report (ReportV2): Report object
        graph_file (str): graph html object

    Returns:
        str: _description_
    """
    sources = set()
    # create sources for the report
    for section in report.report.sections:
        sources.update(section.sources)
    sorted_sources = sorted(sources)
    # join together sentences
    paragraphs = []
    for section in report.report.sections:
        for paragraph in section.paragraphs:
            paragraphs.append(" ".join(paragraph.sentences))
    # templates
    template_paths = [os.path.join(BASEDIR, "templates")]
    if graph_file:
        template_paths.append(os.path.dirname(graph_file))
    env = Environment(loader=FileSystemLoader(template_paths))
    report_template = env.get_template("report_v2.html")
    output_from_parsed_template = report_template.render(
        report_title=report.report.title,
        report_description=report.report.description,
        report_sections=[section.dict() for section in report.report.sections],
        report_sources=sorted_sources,
        graph_file=os.path.basename(graph_file),
    )
    return html.unescape(output_from_parsed_template)


@traceable
def string_to_report(
    string: str,
) -> ParsedReport:
    """
    Generate a report object from a string
    """
    report_parser = PydanticOutputParser(pydantic_object=ParsedReport)
    string_to_report_prompt = PromptTemplate(
        input_variables=["string"],
        template=dedent(
            """
        Convert the following string to a Report JSON object.
        Do not modify the data, just parse it into the JSON object.
        Sections should be marked by a number and a title.
        The key sections are Overview, Market Analysis, SWOT Analysis, Competitors, and Sources.

        Each Section should be be broken down into paragraphs.
        Include newlines in individual sections to that the sections maintain the original document structure when parsed into JSON.
        Always include all the sections and paragraphs from the original report.

        #### Example:
        1. Overview
        Background:
        Acme Corp is a company that specializes in making widgets.

        Acme was founded in 2008.

        Key Personnel:
        - John Doe is the CEO of Acme Corp.
        - Jim Dow is a CTO of Acme Corp.

        Pricing:
        Acme Corp's pricing is competitive.
        They also list their pricing publically.
        Products/Services:
        Acme Corp makes widgets.

        2. Market Analysis
        Market:
        The market for widgets is growing.

        TAM/SAM/SOM:
        The TAM for widgets is $1 billion.

        3. SWOT Analysis
        Strengths:
        1. Great product
        2. Strong Team

        Weaknesses:
        1. Bad location
        2. CEO low morale

        Opportunities:
        1. Cyber data
        2. Targeting new markets

        Threats:
        1. Strong competition
        2. Economic downturn

        4. Competitors
        Competitor 1
        Competitor 1 is a competitor of Acme Corp.

        Competitor 2
        Competitor 2 is a competitor of Acme Corp.

        5. Sources
        - https://acmecorp.com
        - https://acmecorp.com/about

        #### Example JSON Object:
        {{
            "title": "Acme Corp Report",
            "description": "This is a report on Acme Corp.",
            "sections": [
                {{
                    "title": "1. Overview",
                    "paragraphs": [
                        {{
                            "title": "Background:"
                            "content": "\\nAcme Corp is a company that specializes in making widgets.\\nAcme was founded in 2008."
                        }},
                        {{
                            "title": "Key Personnel:",
                            "content": "\\n- John Doe is the CEO of Acme Corp.\\n- Jim Dow is a CTO of Acme Corp."
                        }},
                        {{
                            "title": "Pricing:",
                            "content": "\\nAcme Corp's pricing is competitive.\\nThey also list their pricing publically."
                        }}
                    ]
                }},
                {{
                    "title": "2. Market Analysis",
                    "paragraphs": [
                        {{
                            "title": "Market:"
                            "content": "\\nThe market for widgets is growing."
                        }},
                        {{
                            "title": "TAM/SAM/SOM:"
                            "content": "\\nThe TAM for widgets is $1 billion."
                        }}
                    ]
                }},
                {{
                    "title": "3. SWOT Analysis",
                    "paragraphs": [
                        {{
                            "title": "Strengths:"
                            "content": "\\n1. Great product\\n2. Strong Team"
                        }},
                        {{
                            "title": "Weaknesses:",
                            "content": "\\n1. Bad location\\n2. CEO low morale"
                        }},
                        {{
                            "title": "Opportunities:",
                            "content": "\\n1. Cyber data\\n2. Targeting new markets"
                        }},
                        {{
                            "title": "Threats:",
                            "content": "\\n1. Strong competition\\n2. Economic downturn"
                        }}
                    ]
                }},
                {{
                    "title": "4. Competitors",
                    "paragraphs": [
                        {{
                            "title": "Competitor 1"
                            "content": "\\nCompetitor 1 is a competitor of Acme Corp."
                        }},
                        {{
                            "title": "Competitor 2"
                            "content": "\\nCompetitor 2 is a competitor of Acme Corp."
                        }}
                    ]
                }},
                {{
                    "title": "5. Sources",
                    "paragraphs": [
                        {{
                            "title": "Links"
                            "content": "\\n- https://acmecorp.com\\n- https://acmecorp.com/about"
                        }}
                    ]
                }}
            ]
        }}
        #### End of example
        {string}
        \n
        {format_instructions}
        """
        ),
        partial_variables={
            "format_instructions": report_parser.get_format_instructions()
        },
    )
    chain = LLMChain(
        llm=openai_gpt_4o,
        prompt=string_to_report_prompt,
    )
    response = chain.invoke({"string": string})
    parsed_result = report_parser.parse(text=response["text"])
    return parsed_result


def report_to_pdf(report: Report, filename: str) -> bytes:
    """
    Convert a report to a PDF
    """
    html = report_to_html(report)
    pdf = pdfkit.from_string(html, filename)
    return pdf


def add_watermark(canvas, doc):
    # Define the watermark image file path

    # Define the position (x, y) where the watermark will be placed
    x = 25
    y = 750

    # Draw the watermark image on the page
    if os.getenv("WATERMARK_IMAGE"):
        canvas.drawImage(
            os.getenv("WATERMARK_IMAGE"), x, y, width=75, height=25, mask="auto"
        )
    else:
        logging.info("No watermark image found, will not add watermark to PDF.")

    # add page number
    page_number_text = f"{doc.page}"

    # Define the position (x, y) where the page number will be placed
    x = letter[0] - 100  # Right side
    y = 20  # Bottom side

    # Set the font and size
    canvas.setFont("Helvetica", 10)

    # Draw the page number on the bottom-right of the page
    canvas.drawString(x, y, page_number_text)


def report_v2_to_pdf(
    report: ReportV2, filename: str, graph_file: str = None, watermark: bool = None
) -> SimpleDocTemplate:
    """Generate a PDF report with a graph embedded in the HTML

    Args:
        report (ReportV2): Report object
        filename (str): file to save
        graph_file (str, optional): graph file to include. Defaults to None.
        watermark_file (str, optional): watermark file to include. Defaults to None.

    Returns:
        bytes: _description_
    """
    sources = set()
    document_elements = []
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    title = Paragraph(report.report.title, styles["Title"])
    # append title to document elements
    document_elements.append(title)
    image_counter = 1
    # Create a custom paragraph style with justified alignment
    justified_paragraph_style = ParagraphStyle(
        name="Justified",
        parent=styles["Normal"],  # Use 'Normal' as the base
        alignment=TA_JUSTIFY,  # Justified alignment
    )
    for section in report.report.sections:
        section_title = Paragraph(section.title, styles["Heading2"])
        document_elements.append(section_title)
        for paragraph in section.paragraphs:
            if paragraph.title:
                paragraph_title = Paragraph(paragraph.title, styles["Heading3"])
                document_elements.append(paragraph_title)
            # check if paragraph has an image
            if paragraph.images:
                image_content = get_image(paragraph.images.results[0].original_url)
                if image_content:
                    try:
                        # save image to a temporary file and add to document elements
                        with tempfile.NamedTemporaryFile() as f:
                            f.write(image_content)
                            image = Image(f.name)
                            # format image for table
                            image.drawHeight = (
                                2 * inch * image.drawHeight / image.drawWidth
                            )
                            image.drawWidth = 2 * inch
                            # Create a caption for the image
                            if "Caption" not in styles:
                                caption_style = ParagraphStyle(
                                    name="Caption",
                                    parent=styles["Normal"],  # Use 'Normal' as the base
                                    fontSize=10,
                                    italic=True,
                                    alignment=1,  # Centered alignment
                                    spaceBefore=6,  # Space before the caption
                                    spaceAfter=6,  # Space after the caption
                                )
                                styles.add(caption_style)
                            if paragraph.images.results[0].caption:
                                caption = Paragraph(
                                    f"Figure {image_counter}: {paragraph.images.results[0].caption}",
                                    styles["Caption"],
                                )
                            else:
                                caption = Paragraph(
                                    f"Figure {image_counter}: {paragraph.images.results[0].title.rstrip(" ... ")}",
                                    styles["Caption"],
                                )
                            # Combine image and caption in a nested table
                            image_table = Table([[image], [caption]])
                            image_table.setStyle(
                                TableStyle(
                                    [
                                        (
                                            "ALIGN",
                                            (0, 0),
                                            (-1, -1),
                                            "CENTER",
                                        ),  # Center the image and caption
                                        (
                                            "TOPPADDING",
                                            (0, 0),
                                            (-1, 0),
                                            0,
                                        ),  # Remove top padding of the image
                                        (
                                            "BOTTOMPADDING",
                                            (0, 1),
                                            (-1, -1),
                                            0,
                                        ),  # Remove bottom padding of the caption
                                    ]
                                )
                            )
                            raw_paragraph_content = " ".join(paragraph.sentences)
                            paragraph_content = Paragraph(
                                raw_paragraph_content, justified_paragraph_style
                            )
                            table_data = [[paragraph_content, image_table]]
                            table = Table(table_data)
                            table.setStyle(
                                TableStyle(
                                    [
                                        (
                                            "VALIGN",
                                            (0, 0),
                                            (-1, -1),
                                            "TOP",
                                        ),  # Align text and image to the top
                                        (
                                            "LEFTPADDING",
                                            (0, 0),
                                            (-1, -1),
                                            0,
                                        ),  # Remove left padding
                                        (
                                            "RIGHTPADDING",
                                            (0, 0),
                                            (-1, -1),
                                            10,
                                        ),  # Add some space between image and text
                                        (
                                            "TOPPADDING",
                                            (0, 0),
                                            (-1, -1),
                                            0,
                                        ),  # Remove top padding
                                        (
                                            "BOTTOMPADDING",
                                            (0, 0),
                                            (-1, -1),
                                            0,
                                        ),  # Remove bottom padding
                                        (
                                            "BACKGROUND",
                                            (0, 0),
                                            (-1, -1),
                                            colors.white,
                                        ),  # Background color if needed
                                    ]
                                )
                            )
                            # Add the table to the content
                            image_counter += 1
                            document_elements.append(table)
                    except UnidentifiedImageError:
                        print(
                            f"Unable to load image from {paragraph.images.results[0].original_url}"
                        )
                        continue
            else:
                # append paragraph content to document elements
                raw_paragraph_content = " ".join(paragraph.sentences)
                # add sources to the document
                if section.sources and len(section.sources) > 0:
                    for source in section.sources:
                        sources.add(source)
                paragraph_content = Paragraph(
                    raw_paragraph_content, justified_paragraph_style
                )
                document_elements.append(paragraph_content)
            # add space between paragraphs
            document_elements.append(Spacer(1, 12))
    # add graph to the document
    if graph_file:
        # add graph title
        document_elements.append(PageBreak())
        graph_title = Paragraph("Entity Graph", styles["Heading2"])
        document_elements.append(graph_title)
        document_elements.append(Image(graph_file, width=450, height=400))
        # add space between graph and sources
        document_elements.append(Spacer(1, 12))
    # append sources to the document
    document_elements.append(PageBreak())
    bullet_style = ListStyle(name="BulletStyle", bulletType="bullet")
    sources_header = Paragraph("Sources", styles["Heading2"])
    document_elements.append(sources_header)
    # add sources to the document
    # convert sources to a list
    list_flowable = []
    for source in sources:
        list_flowable.append(
            ListItem(Paragraph(source, styles["BodyText"]), bulletText="•")
        )
    sources_list = ListFlowable(list_flowable, style=bullet_style, bulletType="bullet")
    document_elements.append(sources_list)
    if watermark:
        doc.build(
            document_elements, onFirstPage=add_watermark, onLaterPages=add_watermark
        )
    else:
        doc.build(document_elements)
    return doc


def report_to_pdf_binary(report: Report) -> bytes:
    """
    Convert a report to a PDF
    """
    html = report_to_html(report)
    with tempfile.NamedTemporaryFile() as f:
        pdfkit.from_string(html, f.name)
        return f.read()


def report_to_docx(report: Report) -> Document:
    """
    Convert a report to a DOCX
    """
    document = Document()
    document.add_heading(report.report.title, level=1)
    for section in report.report.sections:
        document.add_heading(section.title, level=2)
        for paragraph in section.paragraphs:
            if paragraph.title and paragraph.title != "":
                document.add_heading(paragraph.title, level=3)
            document.add_paragraph(paragraph.content)
    return document


def report_v2_to_docx(report: ReportV2) -> Document:
    """
    Convert a report to a DOCX
    """
    sources = set()
    # create sources for the report
    for section in report.report.sections:
        sources.update(section.sources)
    sorted_sources = sorted(sources)
    document = Document()
    document.add_heading(report.report.title, level=1)
    for section in report.report.sections:
        document.add_heading(section.title, level=2)
        for paragraph in section.paragraphs:
            if paragraph.title and paragraph.title != "":
                document.add_heading(paragraph.title, level=3)
            document.add_paragraph(" ".join(paragraph.sentences))
    document.add_heading("Sources", level=2)
    for source in sorted_sources:
        document.add_paragraph(source, style="ListBullet")
    return document
