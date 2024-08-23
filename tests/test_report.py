from tests.constants import (
    TEST_REPORT_RESPONSE,
    REPORT_JSON,
    TEST_COMPLEX_NARRATIVE,
    TEST_KEY_QUESTIONS_BULLETED,
    TEST_KEY_QUESTIONS_NARRATIVE,
    CREW_RUN,
    REPORT_V2_JSON,
    GRAPH_PNG,
    BASEDIR,
)
from conductor.reports.outputs import (
    string_to_report,
)
from conductor.reports.models import (
    Report,
    Paragraph,
    Section,
    ParsedReport,
    ReportStyleV2,
    ReportTone,
    ReportPointOfView,
    ReportV2,
)
from conductor.reports.outputs import (
    report_to_docx,
    report_to_html,
    report_v2_to_html,
    report_v2_to_docx,
    report_v2_to_pdf,
)
from conductor.crews.rag_marketing.chains import crew_run_to_report
from conductor.crews.models import CrewRun
from langsmith import unit
from docx.document import Document as DocumentObject
import os


def test_conductor_url_report_generator():
    urls = ["trssllc.com"]
    report_title = "Test Report"
    report_description = "This is a test report."
    section_title = "This is a test section title"
    report = Report(
        title=report_title,
        description=report_description,
        sections=[
            Section(
                title=section_title,
                paragraphs=[
                    Paragraph(
                        title=f"Website Summary of {urls[0]}",
                        content="This is a summary of the website.",
                    ),
                    Paragraph(
                        title=f"Personnel Summary of {urls[0]}",
                        content="This is a summary of the personnel.",
                    ),
                ],
            )
        ],
    )
    assert report.title == report_title
    assert report.description == report_description
    assert len(report.sections) == 1
    assert report.sections[0].title == section_title
    assert len(report.sections[0].paragraphs) == 2
    assert report.sections[0].paragraphs[0].title == f"Website Summary of {urls[0]}"
    assert report.sections[0].paragraphs[1].title == f"Personnel Summary of {urls[0]}"
    assert report.sections[0].paragraphs[0].content is not None
    assert report.sections[0].paragraphs[1].content is not None


@unit
def test_string_to_report() -> None:
    report = string_to_report(
        TEST_REPORT_RESPONSE,
    )
    assert isinstance(report, ParsedReport)
    # assert report.title == "Thomson Reuters Special Services (TRSS) Company Report"
    # # test for section structure
    # assert len(report.sections) == 5
    # # Overview Section
    # assert report.sections[0].title == "1. Overview"
    # assert len(report.sections[0].paragraphs) == 4
    # # Market Analysis Section
    # assert report.sections[1].title == "2. Market Analysis"
    # assert len(report.sections[1].paragraphs) == 2
    # # SWOT Analysis Section
    # assert report.sections[2].title == "3. SWOT Analysis"
    # assert len(report.sections[2].paragraphs) == 4
    # # Competitors Section
    # assert report.sections[3].title == "4. Competitors"
    # assert len(report.sections[3].paragraphs) == 2
    # assert report.raw == TEST_REPORT_RESPONSE
    # # Sources Section
    # assert report.sections[4].title == "5. Sources"
    # assert len(report.sections[4].paragraphs) == 1


def test_report_to_html() -> None:
    # read in the json data
    parsed_report = ParsedReport(**REPORT_JSON)
    report = Report(report=parsed_report)
    html = report_to_html(report)
    assert isinstance(html, str)


def test_complex_narrative_to_report() -> None:
    report = string_to_report(
        TEST_COMPLEX_NARRATIVE,
    )
    assert isinstance(report, ParsedReport)


def test_report_to_docx() -> None:
    # read in the json data
    parsed_report = ParsedReport(**REPORT_JSON)
    report = Report(report=parsed_report)
    document = report_to_docx(report)
    assert isinstance(document, DocumentObject)


def test_key_questions_bulleted_to_report() -> None:
    report = string_to_report(
        TEST_KEY_QUESTIONS_BULLETED,
    )
    assert isinstance(report, ParsedReport)


def test_key_questions_narrative_to_report() -> None:
    report = string_to_report(
        TEST_KEY_QUESTIONS_NARRATIVE,
    )
    assert isinstance(report, ParsedReport)


def test_crew_run_to_report() -> None:
    crew_run = CrewRun.parse_obj(CREW_RUN)
    report = crew_run_to_report(
        crew_run=crew_run,
        title="TRSS Report",
        description="Evrim Insights on TRSS",
        section_titles_endswith_filter="Research",
        tone=ReportTone.INFORMAL,
        style=ReportStyleV2.NARRATIVE,
        point_of_view=ReportPointOfView.THIRD_PERSON,
    )
    assert isinstance(report, ReportV2)


def test_report_v2_to_html() -> None:
    report_v2 = ReportV2.parse_obj(REPORT_V2_JSON)
    html = report_v2_to_html(report=report_v2)
    assert isinstance(html, str)
    # report_v2_to_pdf(report=report_v2, filename="tests/test_report_v2.pdf")


def test_report_v2_to_docx() -> None:
    report_v2 = ReportV2.parse_obj(REPORT_V2_JSON)
    doc = report_v2_to_docx(report=report_v2)
    assert isinstance(doc, DocumentObject)
    # doc.save("tests/test_report_v2.docx")


def test_report_v2_to_pdf() -> None:
    report_v2 = ReportV2.parse_obj(REPORT_V2_JSON)
    report_v2_to_pdf(
        report=report_v2,
        filename=os.path.join(BASEDIR, "data", "test_report_v2.pdf"),
        graph_file=GRAPH_PNG,
        watermark=True,
    )
    assert True
